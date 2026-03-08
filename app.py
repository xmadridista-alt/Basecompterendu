import streamlit as st
import os
import re
from docx import Document
from io import BytesIO

# --- CONFIGURATION ---
st.set_page_config(page_title="DocFetcher Radiologie - Dr Kabaou", layout="wide", page_icon="⚡")

BASE_DIR = "BIBLIOTHEQUE_CR"
CATEGORIES = ["ECHO", "TDM", "IRM", "RX_STANDARD", "DIVERS"]

# --- STYLES CSS (Pour le surlignage Jaune DocFetcher) ---
st.markdown("""
<style>
    mark {
        background-color: #ffff00;
        color: black;
        font-weight: bold;
        padding: 0 2px;
    }
    .cr-viewer {
        background-color: #f9f9f9;
        border: 1px solid #ddd;
        padding: 20px;
        border-radius: 5px;
        font-family: Arial, sans-serif;
        line-height: 1.6;
        height: 600px;
        overflow-y: scroll;
    }
</style>
""", unsafe_allow_html=True)

# --- FONCTIONS ---

@st.cache_data
def indexer_tout(dossier_racine):
    """Lit tous les fichiers et les charge en mémoire"""
    db = []
    if not os.path.exists(dossier_racine):
        os.makedirs(dossier_racine)
        for c in CATEGORIES: os.makedirs(os.path.join(dossier_racine, c))
        return []

    for cat in CATEGORIES:
        path = os.path.join(dossier_racine, cat)
        if os.path.exists(path):
            fichiers = [f for f in os.listdir(path) if f.endswith(".docx") and not f.startswith("~$")]
            for f in fichiers:
                try:
                    doc = Document(os.path.join(path, f))
                    txt = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                    db.append({"nom": f, "cat": cat, "contenu": txt, "chemin": os.path.join(path, f)})
                except: continue
    return db

def surligner_texte(texte, mot_cle):
    """Met en jaune le mot clé en gardant la casse originale (insensible à la casse)"""
    if not mot_cle: return texte.replace("\n", "<br>")
    
    # Regex pour trouver le mot (ignorer majuscules/minuscules) et le wrapper dans <mark>
    # Le pattern ( ... ) permet de garder le texte original trouvé (ex: kyste ou Kyste)
    pattern = re.compile(f"({re.escape(mot_cle)})", re.IGNORECASE)
    texte_surligne = pattern.sub(r"<mark>\1</mark>", texte)
    
    # On remplace les retours à la ligne par des <br> pour l'affichage HTML
    return texte_surligne.replace("\n", "<br>")

def sauvegarder_nouveau(cat, titre, contenu):
    doc = Document()
    doc.add_heading(titre, 0)
    for l in contenu.split('\n'): doc.add_paragraph(l)
    doc.save(os.path.join(BASE_DIR, cat, titre.replace("/", "-") + ".docx"))

# --- LOGIQUE DE SESSION (Pour les boutons Suivant/Précédent) ---
if 'index_actuel' not in st.session_state:
    st.session_state.index_actuel = 0

def page_suivante():
    st.session_state.index_actuel += 1
def page_precedente():
    st.session_state.index_actuel -= 1
def reset_index():
    st.session_state.index_actuel = 0

# --- INTERFACE ---

# Sidebar : Dr Kabaou & Recherche
with st.sidebar:
    st.markdown("## 👨‍⚕️ Dr Kabaou Nadim")
    st.caption("📧 nadimkabaou@gmail.com | 📱 +216 27 561 861")
    st.divider()
    
    st.header("🔍 Filtres")
    choix_cat = st.radio("Modalité :", ["TOUT"] + CATEGORIES, on_change=reset_index)
    
    # Recherche principale
    # Note : on_change=reset_index remet le compteur à 0 si on change le mot clé
    query = st.text_input("Mot-clé à surligner :", placeholder="Ex: Appendicite...", on_change=reset_index)
    
    st.divider()
    if st.button("🔄 Recharger la base"):
        st.cache_data.clear()
        st.rerun()

# Chargement DB
db = indexer_tout(BASE_DIR)

# FILTRAGE DES RÉSULTATS
resultats = []
for item in db:
    match_cat = (choix_cat == "TOUT") or (item['cat'] == choix_cat)
    match_txt = True
    if query:
        # On cherche dans le titre OU le contenu
        if (query.lower() not in item['nom'].lower()) and (query.lower() not in item['contenu'].lower()):
            match_txt = False
    
    if match_cat and match_txt:
        resultats.append(item)

# GESTION DES INDICES (Boutons Next/Prev)
total_res = len(resultats)
if st.session_state.index_actuel >= total_res: st.session_state.index_actuel = 0
if st.session_state.index_actuel < 0: st.session_state.index_actuel = total_res - 1

# --- AFFICHAGE PRINCIPAL ---

st.title("⚡ Visionneuse Rapide")

if total_res == 0:
    st.warning("Aucun résultat trouvé avec ces critères.")
else:
    # 1. BARRE DE NAVIGATION (Top)
    col_prev, col_info, col_next = st.columns([1, 4, 1])
    
    with col_prev:
        if st.button("⬅️ Précédent", use_container_width=True):
            page_precedente()
            st.rerun()
            
    with col_next:
        if st.button("Suivant ➡️", use_container_width=True):
            page_suivante()
            st.rerun()
            
    with col_info:
        # Affichage "Résultat 1 sur 15"
        current = resultats[st.session_state.index_actuel]
        st.markdown(f"<h3 style='text-align: center; margin:0;'>{st.session_state.index_actuel + 1} / {total_res}</h3>", unsafe_allow_html=True)
        st.markdown(f"<div style='text-align: center; color:gray;'>📂 {current['cat']} | 📄 {current['nom']}</div>", unsafe_allow_html=True)

    st.markdown("---")

    # 2. ZONE DE LECTURE (Surlignage)
    
    tab_lecture, tab_copie, tab_ajout = st.tabs(["👁️ Lecture (Surligné)", "📝 Copier / Modifier", "➕ Nouveau"])
    
    with tab_lecture:
        # C'est ici que la magie opère : On injecte du HTML
        contenu_html = surligner_texte(current['contenu'], query)
        
        # On crée une fausse "page" blanche avec le texte
        st.markdown(f"""
        <div class="cr-viewer">
            {contenu_html}
        </div>
        """, unsafe_allow_html=True)

    with tab_copie:
        st.info("Ici, le texte est brut pour faciliter le copier-coller.")
        texte_final = st.text_area("Texte du CR :", value=current['contenu'], height=500)
        st.code(texte_final, language="text")

    with tab_ajout:
        st.write("Ajouter un nouveau CR rapidement :")
        with st.form("ajout_rapide"):
            t_titre = st.text_input("Titre")
            t_cat = st.selectbox("Catégorie", CATEGORIES)
            t_cont = st.text_area("Contenu")
            if st.form_submit_button("Sauvegarder"):
                sauvegarder_nouveau(t_cat, t_titre, t_cont)
                st.success("Enregistré ! Rechargez la page.")
                st.cache_data.clear()